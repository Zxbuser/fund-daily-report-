# =============================================================================
# SECTION A: Imports, Constants, Configuration
# =============================================================================
import concurrent.futures
import datetime as dt
import html
import json
import os
import re
import smtplib
import ssl
import statistics
import subprocess
import sys
import traceback
import urllib.error
import urllib.request
import winreg
from collections import Counter
from email.message import EmailMessage
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --- Paths ---
ROOT = Path(__file__).resolve().parent
REPORT_ROOT = ROOT / "reports"
LOG_DIR = REPORT_ROOT / "logs"

# --- Fund portfolios ---
FUND_ONE = [
    "027052", "021528", "021485", "022365", "026376", "026733",
    "968044", "008971", "021143", "000218", "040046",
]
FUND_TWO = [
    "021528", "022365", "014915", "025209", "011452", "024975",
    "005359", "011892", "018957",
]

# --- Watchlist funds (not yet held, under consideration for entry) ---
WATCHLIST_CODES = ["016370", "006503"]

# --- Email config ---
# Recipients are read from env vars FUND_REPORT_EMAIL_ONE / FUND_REPORT_EMAIL_TWO
# Falls back to placeholder values — edit the env vars with your actual addresses.
RECIPIENTS = {
    "基金一": user_env("FUND_REPORT_EMAIL_ONE") or "your_email_1@example.com",
    "基金二": user_env("FUND_REPORT_EMAIL_TWO") or "your_email_2@example.com",
}
SMTP_HOST = "smtp.gmail.com"
SMTP_PORT = 587

# --- AI API Configuration ---
# Environment variables (read via user_env, which checks os.environ then Windows registry):
#   FUND_REPORT_OPENAI_KEY    - OpenAI API key for ChatGPT
#   FUND_REPORT_ANTHROPIC_KEY - Anthropic API key for Claude
#   FUND_REPORT_DEEPSEEK_KEY  - DeepSeek API key
AI_APIS = {
    "chatgpt": {
        "url": "https://api.openai.com/v1/chat/completions",
        "model": "gpt-4o",
        "key_env": "FUND_REPORT_OPENAI_KEY",
        "timeout": 120,
    },
    "claude": {
        "url": "https://api.anthropic.com/v1/messages",
        "model": "claude-sonnet-4-20250514",
        "key_env": "FUND_REPORT_ANTHROPIC_KEY",
        "timeout": 120,
        "version": "2023-06-01",
    },
    "deepseek": {
        "url": "https://api.deepseek.com/v1/chat/completions",
        "model": "deepseek-chat",
        "key_env": "FUND_REPORT_DEEPSEEK_KEY",
        "timeout": 120,
    },
}

# =============================================================================
# SECTION B: Utility Functions
# =============================================================================


def beijing_now():
    return dt.datetime.now(dt.timezone.utc).replace(tzinfo=None) + dt.timedelta(hours=8)


RUN_DATE = beijing_now().date()
ANALYSIS_DATE = RUN_DATE
RUN_DIR = REPORT_ROOT / RUN_DATE.isoformat()
LOG_FILE = LOG_DIR / f"fund_daily_report_{RUN_DATE.isoformat()}.log"


def log(message):
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    line = f"[{beijing_now().strftime('%Y-%m-%d %H:%M:%S')}] {message}"
    print(line)
    with LOG_FILE.open("a", encoding="utf-8") as fh:
        fh.write(line + "\n")


def user_env(name):
    value = os.environ.get(name)
    if value:
        return value
    try:
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, "Environment") as key:
            value, _ = winreg.QueryValueEx(key, name)
            return value
    except OSError:
        return ""


def curl_text(url):
    raw = subprocess.check_output(
        ["curl.exe", "-s", "-L", "--max-time", "30", url],
        stderr=subprocess.DEVNULL,
    )
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("gb18030", "replace")


# =============================================================================
# SECTION C: Fund Data Parsing (existing + new extractors)
# =============================================================================


def parse_var(text, name):
    match = re.search(rf'var\s+{re.escape(name)}\s*=\s*"(.*?)";', text)
    return match.group(1) if match else ""


def parse_assignment_json(text, var_name):
    match = re.search(rf"var\s+{re.escape(var_name)}\s*=\s*(.*?);/\*", text, re.S)
    if not match:
        match = re.search(rf"var\s+{re.escape(var_name)}\s*=\s*(.*?);", text, re.S)
    if not match:
        return None
    try:
        return json.loads(match.group(1))
    except Exception:
        return None


# --- Existing extractors (unchanged) ---


def fund_latest_networth(text):
    data = parse_assignment_json(text, "Data_netWorthTrend")
    if not data:
        return None
    latest = None
    for item in data:
        day = (
            dt.datetime.fromtimestamp(item["x"] / 1000, dt.timezone.utc).replace(
                tzinfo=None
            )
            + dt.timedelta(hours=8)
        ).date()
        if day <= ANALYSIS_DATE:
            latest = {
                "date": day.isoformat(),
                "nav": item.get("y"),
                "day_return": item.get("equityReturn"),
            }
    return latest


def manager_names(text):
    data = parse_assignment_json(text, "Data_currentFundManager")
    if not data:
        return "--"
    names = [item.get("name") for item in data if item.get("name")]
    return "、".join(names) if names else "--"


def latest_scale(text):
    data = parse_assignment_json(text, "Data_fluctuationScale")
    if not data:
        return "--"
    categories = data.get("categories") or []
    series = data.get("series") or []
    if not categories or not series:
        return "--"
    item = series[-1]
    value = item.get("y")
    mom = item.get("mom")
    if value is None:
        return "--"
    return f"{categories[-1]}：{value}亿" + (f"，环比{mom}" if mom else "")


def latest_allocation(text):
    data = parse_assignment_json(text, "Data_assetAllocation")
    if not data:
        return "--"
    parts = []
    for item in data.get("series", []):
        values = item.get("data") or []
        if values:
            parts.append(f"{item.get('name', '')}{values[-1]}%")
    return "；".join(parts[:4]) if parts else "--"


# --- NEW data extractors for deep analysis ---


def extract_kline_data(text, days=30):
    """Extract daily NAV history for K-line / trend analysis.
    Returns list of {date, nav, day_return} for the most recent `days` entries.
    """
    data = parse_assignment_json(text, "Data_netWorthTrend")
    if not data:
        return []
    result = []
    for item in data:
        day_date = (
            dt.datetime.fromtimestamp(item["x"] / 1000, dt.timezone.utc).replace(
                tzinfo=None
            )
            + dt.timedelta(hours=8)
        ).date()
        if day_date <= ANALYSIS_DATE:
            result.append(
                {
                    "date": day_date.isoformat(),
                    "nav": item.get("y"),
                    "day_return": item.get("equityReturn"),
                }
            )
    return result[-days:] if len(result) > days else result


def extract_kline_summary(kline_data):
    """Compute summary statistics from K-line data for AI prompt context."""
    if not kline_data or len(kline_data) < 3:
        return {
            "period_days": len(kline_data),
            "start_nav": "--",
            "end_nav": "--",
            "period_return": "--",
            "high": "--",
            "low": "--",
            "volatility": "--",
            "up_days": 0,
            "down_days": 0,
            "max_single_day_gain": "--",
            "max_single_day_loss": "--",
            "recent_5d_return": "--",
            "recent_trend": "数据不足",
        }

    navs = [item["nav"] for item in kline_data if item.get("nav") is not None]
    returns = [
        item["day_return"]
        for item in kline_data
        if item.get("day_return") is not None
    ]

    if not navs:
        return {
            "period_days": len(kline_data),
            "recent_trend": "数据不足",
        }

    # Period return
    start_nav = navs[0]
    end_nav = navs[-1]
    period_return = (
        f"{(end_nav / start_nav - 1) * 100:+.2f}%" if start_nav and start_nav != 0 else "--"
    )

    # High / Low
    high_nav = max(navs)
    low_nav = min(navs)

    # Volatility (std of daily returns)
    numeric_returns = [r for r in returns if r is not None]
    if len(numeric_returns) >= 3:
        vol = statistics.stdev(numeric_returns) if len(numeric_returns) > 1 else 0
        volatility = f"{vol:.4f}"
    else:
        volatility = "--"

    # Up/Down days
    up_days = sum(1 for r in numeric_returns if r > 0)
    down_days = sum(1 for r in numeric_returns if r < 0)

    # Max single-day gain/loss
    max_gain = f"{max(numeric_returns):+.2f}%" if numeric_returns else "--"
    max_loss = f"{min(numeric_returns):+.2f}%" if numeric_returns else "--"

    # Recent 5-day return
    recent_5 = navs[-5:] if len(navs) >= 5 else navs
    if len(recent_5) >= 2 and recent_5[0] and recent_5[0] != 0:
        recent_5d_return = f"{(recent_5[-1] / recent_5[0] - 1) * 100:+.2f}%"
    else:
        recent_5d_return = "--"

    # Trend determination using 5-day vs 20-day simple comparison
    if len(navs) >= 5:
        ma5 = sum(navs[-5:]) / 5
        if len(navs) >= 10:
            recent_ma5 = sum(navs[-5:]) / 5
            older_ma5 = sum(navs[-10:-5]) / 5 if len(navs) >= 10 else ma5
            if recent_ma5 > older_ma5 * 1.01:
                recent_trend = "上升"
            elif recent_ma5 < older_ma5 * 0.99:
                recent_trend = "下降"
            else:
                recent_trend = "震荡"
        else:
            recent_trend = "数据有限"
    else:
        recent_trend = "数据不足"

    return {
        "period_days": len(kline_data),
        "start_nav": start_nav,
        "end_nav": end_nav,
        "period_return": period_return,
        "high": f"{high_nav:.4f}",
        "low": f"{low_nav:.4f}",
        "volatility": volatility,
        "up_days": up_days,
        "down_days": down_days,
        "max_single_day_gain": max_gain,
        "max_single_day_loss": max_loss,
        "recent_5d_return": recent_5d_return,
        "recent_trend": recent_trend,
    }


def extract_manager_history(text):
    """Parse Data_managerHistory for manager change records."""
    data = parse_assignment_json(text, "Data_managerHistory")
    if not data:
        return []
    result = []
    for item in data:
        result.append(
            {
                "name": item.get("name", "--"),
                "start_date": item.get("startDate", "--"),
                "end_date": item.get("endDate", "--"),
                "days_in_charge": item.get("days", 0),
                "return_during_tenure": item.get("yield", "--"),
            }
        )
    return result


def extract_industry_allocation(text):
    """Parse Data_industryAllocation for sector weightings.
    Returns top 8 sectors sorted by allocation descending.
    """
    # Try Data_industryAllocation first (newer format)
    data = parse_assignment_json(text, "Data_industryAllocation")
    if not data:
        # Fallback: try Data_assetAllocation series for stock/bond/cash breakdown
        return []

    result = []
    # Data_industryAllocation structure varies; try common patterns
    if isinstance(data, list):
        for item in data:
            name = item.get("name", "") or item.get("industryName", "") or item.get("sector", "")
            ratio = item.get("y", 0) or item.get("ratio", 0) or item.get("percentage", 0)
            if name and ratio:
                result.append({"sector": str(name), "ratio": float(ratio)})
    elif isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, (int, float)) and value > 0:
                result.append({"sector": str(key), "ratio": float(value)})

    result.sort(key=lambda x: x["ratio"], reverse=True)
    return result[:8]


def extract_top_holdings(text):
    """Parse Data_stockHolds for top 10 stock holdings."""
    holdings = parse_assignment_json(text, "Data_stockHolds")
    if not holdings:
        return []
    result = []
    for item in holdings[:10]:
        result.append(
            {
                "code": str(item.get("code", "--")),
                "name": str(item.get("name", "--")),
                "ratio": item.get("ratio", 0) or item.get("y", 0) or 0,
            }
        )
    return result


def extract_additional_returns(text):
    """Extract additional return periods not currently captured."""
    return {
        "ytd": parse_var(text, "syl_thisYear") or "--",
        "2y": parse_var(text, "syl_2n") or "--",
        "3y": parse_var(text, "syl_3n") or "--",
        "5y": parse_var(text, "syl_5n") or "--",
        "since_inception": parse_var(text, "syl_zfn") or "--",
    }


# =============================================================================
# SECTION D: Index & Fund Fetching
# =============================================================================


def fetch_fund(code):
    url = (
        "https://fund.eastmoney.com/pingzhongdata/"
        f"{code}.js?v={int(dt.datetime.now(dt.timezone.utc).timestamp())}"
    )
    try:
        text = curl_text(url)
        name = parse_var(text, "fS_name")
        if not name:
            raise ValueError("fund name not found")
        latest = fund_latest_networth(text)
        kline = extract_kline_data(text, days=30)
        return {
            "code": code,
            "name": name,
            "latest": latest,
            "m1": parse_var(text, "syl_1y") or "--",
            "m3": parse_var(text, "syl_3y") or "--",
            "m6": parse_var(text, "syl_6y") or "--",
            "y1": parse_var(text, "syl_1n") or "--",
            "manager": manager_names(text),
            "manager_history": extract_manager_history(text),
            "scale": latest_scale(text),
            "allocation": latest_allocation(text),
            "industry_allocation": extract_industry_allocation(text),
            "top_holdings": extract_top_holdings(text),
            "kline_summary": extract_kline_summary(kline),
            "additional_returns": extract_additional_returns(text),
            "source": f"https://fund.eastmoney.com/{code}.html",
            "error": "",
        }
    except Exception as exc:
        return {
            "code": code,
            "name": "未检索到",
            "latest": None,
            "m1": "--",
            "m3": "--",
            "m6": "--",
            "y1": "--",
            "manager": "--",
            "manager_history": [],
            "scale": "--",
            "allocation": "--",
            "industry_allocation": [],
            "top_holdings": [],
            "kline_summary": {},
            "additional_returns": {},
            "source": f"https://fund.eastmoney.com/{code}.html",
            "error": str(exc),
        }


def fetch_index(name, secid):
    end = ANALYSIS_DATE.strftime("%Y%m%d")
    beg = (ANALYSIS_DATE - dt.timedelta(days=10)).strftime("%Y%m%d")
    url = (
        "https://push2his.eastmoney.com/api/qt/stock/kline/get?"
        f"secid={secid}&fields1=f1,f2,f3,f4,f5,f6&"
        "fields2=f51,f52,f53,f54,f55,f56,f57,f58,f59,f60,f61&"
        f"klt=101&fqt=1&beg={beg}&end={end}"
    )
    try:
        data = json.loads(curl_text(url)).get("data") or {}
        klines = data.get("klines") or []
        if not klines:
            raise ValueError("no index data")
        parts = klines[-1].split(",")
        return {
            "name": name,
            "date": parts[0],
            "open": parts[1],
            "close": parts[2],
            "high": parts[3],
            "low": parts[4],
            "amount": parts[6],
            "amp": parts[7],
            "pct": parts[8],
            "chg": parts[9],
            "source": "东方财富行情接口",
        }
    except Exception as exc:
        return {
            "name": name,
            "date": "--",
            "close": "--",
            "pct": "--",
            "chg": "--",
            "low": "--",
            "high": "--",
            "source": str(exc),
        }


def fetch_global_markets():
    """Fetch US and global market data from Sina Finance for macro context analysis.
    Returns a dict with US indices and aggregated macro signals.
    """
    global_config = [
        ("纳斯达克综合", "gb_ixic"),
        ("道琼斯工业", "gb_dji"),
    ]
    results = {}
    for name, symbol in global_config:
        url = f"https://hq.sinajs.cn/list={symbol}"
        try:
            req = urllib.request.Request(
                url,
                headers={"Referer": "https://finance.sina.com.cn", "User-Agent": "Mozilla/5.0"},
            )
            ctx = ssl.create_default_context()
            with urllib.request.urlopen(req, timeout=15, context=ctx) as resp:
                text = resp.read().decode("gbk", errors="replace")
            # Parse Sina format: var hq_str_gb_ixic="name,price,pct,date,time,chg,open,high,low,..."
            parts_text = text.split('"')
            if len(parts_text) < 2:
                results[name] = {"name": name, "close": "--", "pct": "--", "error": "parse failed"}
                continue
            fields = parts_text[1].split(",")
            if len(fields) < 9:
                results[name] = {"name": name, "close": "--", "pct": "--", "error": "insufficient fields"}
                continue
            results[name] = {
                "name": fields[0] or name,
                "close": fields[1],
                "pct": fields[2],
                "chg": fields[4],
                "source": "新浪财经全球行情",
            }
        except Exception as exc:
            results[name] = {"name": name, "close": "--", "pct": "--", "error": str(exc)}
    # Compute macro signals for AI context
    signals = []
    for name, data in results.items():
        if data.get("close") and data["close"] != "--":
            pct_val = to_float(data.get("pct"))
            trend = (
                "上涨" if pct_val is not None and pct_val > 0
                else ("下跌" if pct_val is not None and pct_val < 0 else "平盘")
            )
            signals.append(
                f"{name}：{data['close']}（{trend} {data['pct'] if pct_val is not None else '--'}）"
            )
    macro_context = "；".join(signals) if signals else "全球市场数据暂不可用"
    return {
        "indices": results,
        "macro_signals": macro_context,
        "us_overall": (
            "risk_on" if _count_positive(results) >= 2
            else ("risk_off" if _count_negative(results) >= 2 else "mixed")
        ),
    }


def _count_positive(results):
    count = 0
    for data in results.values():
        pct_val = to_float(data.get("pct"))
        if pct_val is not None and pct_val > 0:
            count += 1
    return count


def _count_negative(results):
    count = 0
    for data in results.values():
        pct_val = to_float(data.get("pct"))
        if pct_val is not None and pct_val < 0:
            count += 1
    return count


def to_float(value):
    try:
        return float(value)
    except Exception:
        return None


def pct_text(value):
    number = to_float(value)
    if number is None:
        return "--"
    return f"{number:+.2f}%"


def group_stats(codes, funds):
    values = []
    for code in codes:
        latest = funds[code].get("latest") or {}
        number = to_float(latest.get("day_return"))
        if number is not None:
            values.append(number)
    if not values:
        return {"count": 0, "avg": "--", "up": 0, "down": 0}
    return {
        "count": len(values),
        "avg": f"{sum(values) / len(values):+.2f}%",
        "up": sum(1 for value in values if value > 0),
        "down": sum(1 for value in values if value < 0),
    }


# =============================================================================
# SECTION E: AI API Integration
# =============================================================================

AI_SYSTEM_PROMPT = """你是一位专业的基金投资分析师，拥有10年以上公募基金研究经验。你需要结合**全球金融市场局势、美股与A股行情、市场情绪以及基金ETF的底层持仓股票**，对提供的基金组合进行全面深入的分析，并给出具体的仓位操作建议。

## 分析框架
对每只基金，请从以下维度进行深度分析并打分（1-10分）：

1. **全球宏观与市场情绪分析**: 结合全球市场信号（美股三大指数涨跌、risk_on/risk_off情绪），判断当前全球资金流向和风险偏好，分析其对A股和港股的影响方向。特别注意：美股科技股走势对A股半导体/科技类基金的传导效应；美元流动性变化对港股通基金的影响。
2. **A股市场环境分析**: 结合上证指数、深证成指、创业板指、科创50、沪深300、恒生指数的当日表现，判断A股整体风格（成长vs价值、大盘vs小盘、科技vs消费），分析当前市场热点和资金偏好。
3. **技术面分析（K线趋势）**: 根据近1个月净值走势判断趋势方向（上升/下降/震荡），识别关键支撑位/压力位，判断短期动量和超买超卖信号。
4. **基金经理评估**: 根据任职时长、任期回报、管理经验评估经理能力和稳定性，结合当前市场风格判断经理是否匹配。
5. **行业配置与持仓分析**: 评估行业集中度、重仓股质量——**特别关注基金ETF的底层持仓股票与当前市场热点的契合度**。比如：重仓半导体是否受益于AI产业链景气？重仓港股是否受外资回流支撑？重仓消费是否与内需复苏匹配？分析重仓股在当前全球局势下的机会与风险。
6. **规模与流动性**: 评估基金规模是否合理（过大影响灵活性，过小有清盘风险），在risk_off情绪下小盘基金是否有流动性风险。
7. **综合评估**: 结合全球局势、市场情绪、技术面、基本面等多维度给出综合评分。

## 仓位建议规则（重要：假设每只基金当前投入为10,000元人民币）
- "增持": 看好后市，建议增加仓位 → 可选百分比: 5%, 10%, 15%, 20%, 25%, 30% （对应500/1000/1500/2000/2500/3000元）
- "减持": 看淡后市或有风险，建议减少仓位 → 可选百分比: 5%, 10%, 15%, 20%, 25%, 30% （对应500/1000/1500/2000/2500/3000元）
- "持有": 维持当前仓位不变，观望为主
- 组合整体应保持审慎：总增持金额不应大幅超过总减持金额
- 对于数据不足的基金（如QDII净值滞后、新基金数据短），倾向建议"持有"
- 风险管理提示：在risk_off（避险）情绪下，应更倾向于减仓或持有高弹性品种

## 输出格式（严格遵守JSON格式，不要添加任何额外说明文字）
```json
{
  "global_macro_analysis": "对全球金融局势和市场情绪的总体判断，结合美股走势、A股风格、资金流向进行分析（3-4句话）",
  "market_overview": "对当前A股/港股市场环境的整体判断（2-3句话）",
  "portfolio_analysis": "对组合整体结构、风险暴露、集中度的分析，特别关注底层持仓股票与全球热点的关联（2-3句话）",
  "funds": [
    {
      "code": "基金代码（6位字符串）",
      "name": "基金名称",
      "technical_score": 7,
      "manager_score": 8,
      "allocation_score": 6,
      "scale_score": 7,
      "overall_score": 7,
      "trend_analysis": "K线趋势分析，包含支撑位/压力位判断（1-2句话）",
      "holding_stock_analysis": "结合该基金ETF的重仓股和当前全球/行业趋势，分析重仓股的机会与风险（1-2句话）",
      "risk_warning": "需要特别关注的风险点（1句话）",
      "action": "增持",
      "percentage_change": 20,
      "amount_change": 2000,
      "reasoning": "操作理由，结合全球局势、市场情绪和基金自身数据说明（1-2句话）"
    }
  ],
  "portfolio_recommendation": "组合层面的整体操作建议和风险提示，结合全球宏观背景给出方向性指引（2-3句话）"
}
```

请严格基于我提供的数据进行分析，不要编造任何信息。如果某只基金数据不足或异常，在reasoning中说明并建议"持有"。

## 特别提醒
- 你的分析必须体现**全球→A股→行业→个股（重仓股）**的传导逻辑
- 如果全球risk_off情绪明显，应提醒关注港股通、QDII等海外暴露基金的波动风险
- 分析重仓股时，不要凭空编造股票代码和名称，只使用我提供的重仓股数据"""

WATCHLIST_SYSTEM_PROMPT = """你是一位专业的基金投资分析师，拥有10年以上公募基金研究经验。你需要结合**全球金融市场局势、美股与A股行情、市场情绪以及基金ETF的底层持仓股票**，对**待买入的自选观察基金**进行深度分析，判断当前是否为合适的买入时机，并给出具体的买入建议。

## 分析框架
对每只基金，请从以下维度进行深度分析并打分（1-10分）：

1. **全球宏观与市场情绪对入场时机的影响**: 结合全球市场信号（美股三大指数涨跌、risk_on/risk_off情绪），判断当前全球资金流向——risk_off情绪下通常不是好的买入时机，risk_on环境下可以更积极寻找入场点。特别关注：美股科技股大幅下跌是否会拖累A股相关基金；全球避险情绪是否会将资金推入黄金/红利等防御品种。
2. **趋势方向与强度**: 根据近期K线走势判断趋势（上升/下降/震荡），分析趋势的强度和可持续性。上升趋势回踩支撑位是最佳买点；下降趋势中即使超跌也不急于抄底。
3. **入场点位评估**: 判断当前净值是否处于合理的买入区间，是否接近支撑位或压力位。结合全球市场波动评估：如果美股隔夜波动剧烈，A股次日开盘可能提供更好的入场点。
4. **基金经理评估**: 根据任职时长、任期回报、管理经验评估经理能力和稳定性，在波动市中经验丰富的经理更具优势。
5. **行业配置与持仓分析**: 评估行业集中度、重仓股质量——**特别关注基金ETF的底层持仓股票当前是否处于上涨趋势中**。分析：重仓股所属行业是否受益于当前全球/国内热点？持仓股估值是否合理？行业是否面临政策或外部风险？
6. **规模与流动性**: 评估基金规模是否合理（过大影响灵活性，过小有清盘风险），在risk_off情绪下小盘基金是否有赎回压力导致的流动性风险。

## 买入建议规则
- **如果当前位置适合买入（can_buy = true）**，请给出：
  - 建议买入金额（suggested_amount，单位：元人民币，范围 1000-20000）
  - 买入策略（entry_strategy）："一次性买入" 或 "分批买入"
  - 具体的买入计划（entry_plan）：包括分几批、每批金额、触发条件
  - 目标价位区间（target_price）
  - 止损位或止损条件（stop_loss）
  - 支撑位（support_level）和压力位（resistance_level）
- **如果当前位置不适合买入（can_buy = false）**，请在 conditions_to_wait_for 中说明需要等待的具体条件（包括全球市场信号方面的条件），suggested_amount 设为 0

## 输出格式（严格遵守JSON格式，不要添加任何额外说明文字）
```json
{
  "global_macro_analysis": "从全球金融局势和美股走势角度，分析当前大环境对新买入决策的影响（2-3句话）",
  "market_overview": "对当前A股/港股市场环境的整体判断，特别是对新买入决策的影响（2-3句话）",
  "funds": [
    {
      "code": "基金代码（6位字符串）",
      "name": "基金名称",
      "can_buy": true,
      "technical_score": 7,
      "manager_score": 8,
      "allocation_score": 6,
      "scale_score": 7,
      "overall_score": 7,
      "trend_analysis": "趋势方向和强度分析，包含短期、中期趋势判断（1-2句话）",
      "entry_analysis": "当前位置是否为良好买点的判断，结合全球市场和A股情绪，包含支撑位/压力位分析（1-2句话）",
      "holding_stock_analysis": "结合该基金ETF的重仓股和当前全球/行业趋势，分析重仓股当前是否处于有利阶段（1-2句话）",
      "suggested_amount": 5000,
      "entry_strategy": "分批买入",
      "entry_plan": "建议分3批买入，首次买入2000元，之后每下跌3%加仓1500元，总投入5000元。若净值突破压力位可追加",
      "support_level": "支撑位净值约1.23（或描述性说明）",
      "resistance_level": "压力位净值约1.35（或描述性说明）",
      "stop_loss": "若净值跌破1.15或从买入点下跌超过8%，建议止损",
      "target_price": "目标区间1.35-1.45",
      "conditions_to_wait_for": "",
      "risk_warning": "需要特别关注的风险点（1句话）",
      "reasoning": "综合研判理由，结合全球局势、市场情绪和数据说明为什么认为可以/不可以买入（1-2句话）"
    }
  ],
  "watchlist_summary": "整体观察建议和组合层面的提示，结合全球宏观背景给出方向性指引（2-3句话）"
}
```

注意：如果 can_buy 为 false，则 suggested_amount 应为 0，entry_strategy 和 entry_plan 可留空，但必须在 conditions_to_wait_for 中说明等待条件（包括全球市场信号方面的等待条件）。

请严格基于我提供的数据进行分析，不要编造任何信息。如果某只基金数据不足或异常，在 reasoning 中说明并建议暂不买入。

## 特别提醒
- 你的分析必须体现**全球→A股→行业→个股（重仓股）**的传导逻辑
- 全球risk_off时通常不是最佳买入时机，建议等待企稳信号；risk_on时积极寻找回调入场点
- 分析重仓股时，不要凭空编造股票代码和名称，只使用我提供的重仓股数据"""


def call_ai_api(api_config, system_prompt, user_prompt):
    """Call an AI API (OpenAI-compatible or Anthropic) and return the response.
    Returns: {"provider": str, "content": str, "error": str|None}
    """
    provider_name = (
        api_config["key_env"]
        .replace("FUND_REPORT_", "")
        .replace("_KEY", "")
        .lower()
    )
    api_key = user_env(api_config["key_env"])
    if not api_key:
        return {
            "provider": provider_name,
            "content": "",
            "error": f"Missing env var: {api_config['key_env']}",
        }

    headers = {}
    body = {}
    is_anthropic = "anthropic" in api_config["url"]

    if is_anthropic:
        headers = {
            "x-api-key": api_key,
            "anthropic-version": api_config.get("version", "2023-06-01"),
            "content-type": "application/json",
        }
        body = {
            "model": api_config["model"],
            "max_tokens": 4096,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        }
    else:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": user_prompt})
        body = {
            "model": api_config["model"],
            "messages": messages,
            "max_tokens": 4096,
            "temperature": 0.3,
        }

    timeout = api_config.get("timeout", 120)
    try:
        req = urllib.request.Request(
            api_config["url"],
            data=json.dumps(body).encode("utf-8"),
            headers=headers,
            method="POST",
        )
        # Allow up to 2 retries for transient errors
        for attempt in range(2):
            try:
                ctx = ssl.create_default_context()
                with urllib.request.urlopen(req, timeout=timeout, context=ctx) as resp:
                    resp_body = json.loads(resp.read().decode("utf-8"))

                if is_anthropic:
                    content = "".join(
                        block.get("text", "")
                        for block in resp_body.get("content", [])
                        if block.get("type") == "text"
                    )
                else:
                    content = (
                        resp_body.get("choices", [{}])[0]
                        .get("message", {})
                        .get("content", "")
                    )

                return {
                    "provider": provider_name,
                    "content": content,
                    "error": None,
                }
            except (urllib.error.HTTPError, urllib.error.URLError,
                    ConnectionResetError, TimeoutError, OSError) as e:
                if attempt == 0:
                    log(f"AI {provider_name} attempt {attempt + 1} failed: {e}, retrying in 5s...")
                    import time as _time
                    _time.sleep(5)
                else:
                    raise

    except urllib.error.HTTPError as e:
        error_body = ""
        try:
            error_body = e.read().decode("utf-8", errors="replace")[:500]
        except Exception:
            pass
        return {
            "provider": provider_name,
            "content": "",
            "error": f"HTTP {e.code}: {error_body}",
        }
    except Exception as e:
        return {"provider": provider_name, "content": "", "error": str(e)}


def build_fund_context_for_ai(label, codes, funds, indices, global_markets=None):
    """Build a structured text block with all fund data for AI analysis."""
    if global_markets is None:
        global_markets = {}
    lines = []
    lines.append(f"=== {label}（{len(codes)}只基金）投资组合分析请求 ===")
    lines.append(f"分析日期：{ANALYSIS_DATE.isoformat()}")
    lines.append(f"数据源：东方财富/天天基金公开数据")
    lines.append(f"假设：每只基金当前投入金额为 10,000 元人民币")
    lines.append("")

    # --- NEW: Global macro context ---
    lines.append("## 全球宏观与市场情绪")
    lines.append(f"  全球市场信号：{global_markets.get('macro_signals', '暂不可用')}")
    lines.append(f"  美国市场情绪：{global_markets.get('us_overall', '暂不可用')}（risk_on=风险偏好，risk_off=避险，mixed=分化）")
    lines.append("")

    # Market overview section
    lines.append("## 一、当前市场环境")
    for idx in indices:
        lines.append(
            f"  {idx['name']}：收盘{idx['close']}，涨跌幅{idx['pct']}%，"
            f"涨跌额{idx['chg']}，区间{idx['low']} - {idx['high']}"
        )
    lines.append("")

    # Per-fund detailed data
    lines.append(f"## 二、{label}持仓基金明细")
    for code in codes:
        f = funds[code]
        latest = f.get("latest") or {}
        ks = f.get("kline_summary", {})
        mgr_history = f.get("manager_history", [])
        ind_alloc = f.get("industry_allocation", [])
        holdings = f.get("top_holdings", [])
        addl = f.get("additional_returns", {})

        lines.append(f"### 基金 {code} —— {f['name']}")
        if f.get("error"):
            lines.append(f"  ⚠ 数据获取异常：{f['error']}")
            lines.append("")
            continue

        # Basic data
        lines.append(f"  - 最新净值：{latest.get('nav', '--')}（日期：{latest.get('date', '--')}）")
        lines.append(f"  - 日涨跌：{pct_text(latest.get('day_return'))}")
        lines.append(
            f"  - 阶段收益：近1月 {pct_text(f['m1'])} ｜ 近3月 {pct_text(f['m3'])} ｜ "
            f"近6月 {pct_text(f['m6'])} ｜ 近1年 {pct_text(f['y1'])}"
        )
        lines.append(
            f"  - 更长周期：今年以来 {pct_text(addl.get('ytd'))} ｜ 近2年 {pct_text(addl.get('2y'))} ｜ "
            f"近3年 {pct_text(addl.get('3y'))} ｜ 近5年 {pct_text(addl.get('5y'))}"
        )

        # Fund manager
        lines.append(f"  - 基金经理：{f['manager']}")
        if mgr_history:
            current = mgr_history[0]
            tenure_years = current.get("days_in_charge", 0) / 365.0
            lines.append(
                f"    现任经理任职约{tenure_years:.1f}年（{current.get('days_in_charge', '--')}天），"
                f"任期内回报：{current.get('return_during_tenure', '--')}"
            )
            if len(mgr_history) > 1:
                past_managers = ", ".join(
                    m["name"] for m in mgr_history[1:4]
                )
                lines.append(f"    历史经理变更{len(mgr_history) - 1}次，历任：{past_managers}")

        # Scale
        lines.append(f"  - 基金规模：{f['scale']}")

        # Asset allocation
        lines.append(f"  - 资产配置：{f['allocation']}")

        # Industry allocation
        if ind_alloc:
            top_sectors = "、".join(
                f"{s['sector']}({s['ratio']:.1f}%)" for s in ind_alloc[:5]
            )
            lines.append(f"  - 行业配置（前5）：{top_sectors}")

        # Top holdings
        if holdings:
            top_stocks = "、".join(
                f"{h['name']}({h['ratio']:.1f}%)" for h in holdings[:5]
            )
            lines.append(f"  - 重仓股（前5）：{top_stocks}")

        # K-line summary
        if ks:
            lines.append(f"  - K线技术数据（近{ks.get('period_days', '--')}个交易日）：")
            lines.append(
                f"    区间收益：{ks.get('period_return', '--')}，"
                f"最高净值：{ks.get('high', '--')}，最低净值：{ks.get('low', '--')}"
            )
            lines.append(
                f"    波动率：{ks.get('volatility', '--')}，"
                f"上涨{ks.get('up_days', 0)}天 / 下跌{ks.get('down_days', 0)}天"
            )
            lines.append(
                f"    近5日收益：{ks.get('recent_5d_return', '--')}，"
                f"趋势判断：{ks.get('recent_trend', '--')}"
            )
            lines.append(
                f"    最大单日涨幅：{ks.get('max_single_day_gain', '--')}，"
                f"最大单日跌幅：{ks.get('max_single_day_loss', '--')}"
            )

        lines.append("")

    lines.append("## 三、分析要求")
    lines.append("请对以上每只基金进行深度分析，并严格按照JSON格式返回结果。")
    lines.append("记住：假设每只基金当前投入10,000元，给出具体的增持/减持百分比和金额。")
    return "\n".join(lines)


def build_watchlist_context_for_ai(codes, funds, indices, global_markets=None):
    """Build a structured text block with watchlist fund data for AI entry analysis."""
    if global_markets is None:
        global_markets = {}
    lines = []
    lines.append(f"=== 自选观察基金（{len(codes)}只）买入分析请求 ===")
    lines.append(f"分析日期：{ANALYSIS_DATE.isoformat()}")
    lines.append(f"数据源：东方财富/天天基金公开数据")
    lines.append("")
    lines.append("注意：以下基金当前**未持有**，本次分析用于判断是否应当买入以及如何买入。")
    lines.append("")

    # --- NEW: Global macro context ---
    lines.append("## 全球宏观与市场情绪")
    lines.append(f"  全球市场信号：{global_markets.get('macro_signals', '暂不可用')}")
    lines.append(f"  美国市场情绪：{global_markets.get('us_overall', '暂不可用')}（risk_on=风险偏好，risk_off=避险，mixed=分化）")
    lines.append("")

    # Market overview section
    lines.append("## 一、当前市场环境")
    for idx in indices:
        lines.append(
            f"  {idx['name']}：收盘{idx['close']}，涨跌幅{idx['pct']}%，"
            f"涨跌额{idx['chg']}，区间{idx['low']} - {idx['high']}"
        )
    lines.append("")

    # Per-fund detailed data (same rendering as build_fund_context_for_ai)
    lines.append(f"## 二、自选观察基金明细")
    for code in codes:
        f = funds[code]
        latest = f.get("latest") or {}
        ks = f.get("kline_summary", {})
        mgr_history = f.get("manager_history", [])
        ind_alloc = f.get("industry_allocation", [])
        holdings = f.get("top_holdings", [])
        addl = f.get("additional_returns", {})

        lines.append(f"### 基金 {code} —— {f['name']}")
        if f.get("error"):
            lines.append(f"  ⚠ 数据获取异常：{f['error']}")
            lines.append("")
            continue

        # Basic data
        lines.append(f"  - 最新净值：{latest.get('nav', '--')}（日期：{latest.get('date', '--')}）")
        lines.append(f"  - 日涨跌：{pct_text(latest.get('day_return'))}")
        lines.append(
            f"  - 阶段收益：近1月 {pct_text(f['m1'])} ｜ 近3月 {pct_text(f['m3'])} ｜ "
            f"近6月 {pct_text(f['m6'])} ｜ 近1年 {pct_text(f['y1'])}"
        )
        lines.append(
            f"  - 更长周期：今年以来 {pct_text(addl.get('ytd'))} ｜ 近2年 {pct_text(addl.get('2y'))} ｜ "
            f"近3年 {pct_text(addl.get('3y'))} ｜ 近5年 {pct_text(addl.get('5y'))}"
        )

        # Fund manager
        lines.append(f"  - 基金经理：{f['manager']}")
        if mgr_history:
            current = mgr_history[0]
            tenure_years = current.get("days_in_charge", 0) / 365.0
            lines.append(
                f"    现任经理任职约{tenure_years:.1f}年（{current.get('days_in_charge', '--')}天），"
                f"任期内回报：{current.get('return_during_tenure', '--')}"
            )
            if len(mgr_history) > 1:
                past_managers = ", ".join(
                    m["name"] for m in mgr_history[1:4]
                )
                lines.append(f"    历史经理变更{len(mgr_history) - 1}次，历任：{past_managers}")

        # Scale
        lines.append(f"  - 基金规模：{f['scale']}")

        # Asset allocation
        lines.append(f"  - 资产配置：{f['allocation']}")

        # Industry allocation
        if ind_alloc:
            top_sectors = "、".join(
                f"{s['sector']}({s['ratio']:.1f}%)" for s in ind_alloc[:5]
            )
            lines.append(f"  - 行业配置（前5）：{top_sectors}")

        # Top holdings
        if holdings:
            top_stocks = "、".join(
                f"{h['name']}({h['ratio']:.1f}%)" for h in holdings[:5]
            )
            lines.append(f"  - 重仓股（前5）：{top_stocks}")

        # K-line summary
        if ks:
            lines.append(f"  - K线技术数据（近{ks.get('period_days', '--')}个交易日）：")
            lines.append(
                f"    区间收益：{ks.get('period_return', '--')}，"
                f"最高净值：{ks.get('high', '--')}，最低净值：{ks.get('low', '--')}"
            )
            lines.append(
                f"    波动率：{ks.get('volatility', '--')}，"
                f"上涨{ks.get('up_days', 0)}天 / 下跌{ks.get('down_days', 0)}天"
            )
            lines.append(
                f"    近5日收益：{ks.get('recent_5d_return', '--')}，"
                f"趋势判断：{ks.get('recent_trend', '--')}"
            )
            lines.append(
                f"    最大单日涨幅：{ks.get('max_single_day_gain', '--')}，"
                f"最大单日跌幅：{ks.get('max_single_day_loss', '--')}"
            )

        lines.append("")

    lines.append("## 三、分析要求")
    lines.append("请对以上每只基金进行买入时机分析，判断当前是否适合买入，并严格按照JSON格式返回结果。")
    lines.append("重点分析：趋势方向、入场点位、支撑/压力位，给出具体的买入金额和策略。")
    lines.append("如果当前不适合买入，请明确说明需要等待的条件。")
    return "\n".join(lines)


def parse_ai_fund_analysis(ai_content, codes):
    """Extract JSON from AI response and parse fund recommendations.
    Returns dict keyed by fund code, or empty dict on failure.
    """
    try:
        # Try markdown code fence extraction first
        json_match = re.search(r"```json\s*(.*?)\s*```", ai_content, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(1))
        else:
            # Try to find JSON object directly
            json_match = re.search(r"\{[\s\S]*\"funds\"[\s\S]*\}", ai_content)
            if json_match:
                data = json.loads(json_match.group(0))
            else:
                data = json.loads(ai_content)

        funds_list = data.get("funds", [])
        result = {}
        for item in funds_list:
            code = str(item.get("code", ""))
            if code in codes:
                result[code] = item

        # Store market_overview and portfolio_recommendation if present
        if data.get("market_overview") or data.get("portfolio_recommendation") or data.get("global_macro_analysis"):
            result["_meta"] = {
                "market_overview": data.get("market_overview", ""),
                "portfolio_analysis": data.get("portfolio_analysis", ""),
                "portfolio_recommendation": data.get("portfolio_recommendation", ""),
                "global_macro_analysis": data.get("global_macro_analysis", ""),
            }

        return result
    except (json.JSONDecodeError, KeyError, TypeError) as e:
        log(f"Failed to parse AI JSON response: {e}")
        return {}


def cross_validate_ai_results(ai_results, codes):
    """Cross-validate AI outputs and produce consensus recommendations.
    - Majority vote for action (增持/减持/持有)
    - Median for percentage change
    - Average for scores
    """
    parsed = {}
    for provider, result in ai_results.items():
        if result.get("content") and not result.get("error"):
            parsed[provider] = parse_ai_fund_analysis(result["content"], set(codes))

    if not parsed:
        return {
            "error": "No AI returned parseable results",
            "funds": {},
            "market_overview": "",
            "portfolio_analysis": "",
            "portfolio_recommendation": "",
            "ai_count": 0,
        }

    # Collect meta info from all AIs
    market_overviews = []
    portfolio_analyses = []
    portfolio_recommendations = []
    global_macro_analyses = []
    for fund_map in parsed.values():
        meta = fund_map.get("_meta", {})
        if meta.get("market_overview"):
            market_overviews.append(f"[{len(market_overviews) + 1}] {meta['market_overview']}")
        if meta.get("portfolio_analysis"):
            portfolio_analyses.append(meta["portfolio_analysis"])
        if meta.get("portfolio_recommendation"):
            portfolio_recommendations.append(meta["portfolio_recommendation"])
        if meta.get("global_macro_analysis"):
            global_macro_analyses.append(meta["global_macro_analysis"])

    consensus_funds = {}
    for code in codes:
        actions = []
        percentages = []
        scores = {
            "technical": [],
            "manager": [],
            "allocation": [],
            "scale": [],
            "overall": [],
        }
        reasonings = []
        trend_analyses = []
        holding_stock_analyses = []
        risk_warnings = []

        for provider, fund_map in parsed.items():
            item = fund_map.get(code)
            if not item:
                continue
            actions.append(item.get("action", "持有"))
            pct = item.get("percentage_change", 0)
            if isinstance(pct, (int, float)) and pct:
                percentages.append(float(pct))
            # Scores
            for key in scores:
                val = item.get(f"{key}_score", 0)
                if isinstance(val, (int, float)) and val:
                    scores[key].append(float(val))
            if item.get("reasoning"):
                reasonings.append(f"[{provider}] {item['reasoning']}")
            if item.get("trend_analysis"):
                trend_analyses.append(item["trend_analysis"])
            if item.get("holding_stock_analysis"):
                holding_stock_analyses.append(item["holding_stock_analysis"])
            if item.get("risk_warning"):
                risk_warnings.append(item["risk_warning"])

        # Majority vote for action
        action_counts = Counter(actions)
        majority_action = (
            action_counts.most_common(1)[0][0] if action_counts else "持有"
        )
        agreement = (
            action_counts[majority_action] / len(actions) if actions else 0
        )

        # Median percentage
        percentages.sort()
        median_pct = statistics.median(percentages) if percentages else 0

        # Average scores (rounded to 1 decimal)
        avg_scores = {}
        for k, v in scores.items():
            avg_scores[f"{k}_score"] = round(sum(v) / len(v), 1) if v else 0

        consensus_funds[code] = {
            "code": code,
            "action": majority_action,
            "ai_agreement": f"{agreement:.0%}",
            "ai_count": len(actions),
            "percentage_change": round(median_pct),
            "amount_change": round(median_pct * 100),  # pct * 10000 CNY / 100
            "technical_score": avg_scores.get("technical_score", 0),
            "manager_score": avg_scores.get("manager_score", 0),
            "allocation_score": avg_scores.get("allocation_score", 0),
            "scale_score": avg_scores.get("scale_score", 0),
            "overall_score": avg_scores.get("overall_score", 0),
            "trend_analysis": trend_analyses[0] if trend_analyses else "--",
            "holding_stock_analysis": holding_stock_analyses[0] if holding_stock_analyses else "--",
            "risk_warning": risk_warnings[0] if risk_warnings else "--",
            "reasoning": " | ".join(reasonings[:3]) if reasonings else "--",
        }

    return {
        "funds": consensus_funds,
        "market_overview": "\n".join(market_overviews) if market_overviews else "",
        "portfolio_analysis": " ".join(portfolio_analyses) if portfolio_analyses else "",
        "portfolio_recommendation": " ".join(portfolio_recommendations) if portfolio_recommendations else "",
        "global_macro_analysis": " ".join(global_macro_analyses) if global_macro_analyses else "",
        "ai_count": len(parsed),
        "error": "",
    }


def parse_watchlist_analysis(ai_content, codes):
    """Extract JSON from AI response and parse watchlist fund recommendations.
    Returns dict keyed by fund code, or empty dict on failure.
    """
    try:
        json_match = re.search(r"```json\s*(.*?)\s*```", ai_content, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group(1))
        else:
            json_match = re.search(r"\{[\s\S]*\"funds\"[\s\S]*\}", ai_content)
            if json_match:
                data = json.loads(json_match.group(0))
            else:
                data = json.loads(ai_content)

        funds_list = data.get("funds", [])
        result = {}
        for item in funds_list:
            code = str(item.get("code", ""))
            if code in codes:
                result[code] = item

        if data.get("market_overview") or data.get("watchlist_summary") or data.get("global_macro_analysis"):
            result["_meta"] = {
                "market_overview": data.get("market_overview", ""),
                "watchlist_summary": data.get("watchlist_summary", ""),
                "global_macro_analysis": data.get("global_macro_analysis", ""),
            }

        return result
    except (json.JSONDecodeError, KeyError, TypeError) as e:
        log(f"Failed to parse watchlist AI JSON response: {e}")
        return {}


def cross_validate_watchlist_results(ai_results, codes):
    """Cross-validate AI outputs and produce consensus for watchlist funds.
    - Majority vote for can_buy (true/false)
    - Median for suggested_amount
    - Average for scores
    - First non-empty text for analysis fields
    """
    parsed = {}
    for provider, result in ai_results.items():
        if result.get("content") and not result.get("error"):
            parsed[provider] = parse_watchlist_analysis(result["content"], set(codes))

    if not parsed:
        return {
            "error": "No AI returned parseable results for watchlist",
            "funds": {},
            "market_overview": "",
            "watchlist_summary": "",
            "ai_count": 0,
        }

    # Collect meta info from all AIs
    market_overviews = []
    watchlist_summaries = []
    global_macro_analyses = []
    for fund_map in parsed.values():
        meta = fund_map.get("_meta", {})
        if meta.get("market_overview"):
            market_overviews.append(meta["market_overview"])
        if meta.get("watchlist_summary"):
            watchlist_summaries.append(meta["watchlist_summary"])
        if meta.get("global_macro_analysis"):
            global_macro_analyses.append(meta["global_macro_analysis"])

    consensus_funds = {}
    for code in codes:
        can_buy_votes = []
        amounts = []
        scores = {
            "technical": [],
            "manager": [],
            "allocation": [],
            "scale": [],
            "overall": [],
        }
        reasonings = []
        trend_analyses = []
        entry_analyses = []
        holding_stock_analyses = []
        risk_warnings = []
        first_text = {}

        for provider, fund_map in parsed.items():
            item = fund_map.get(code)
            if not item:
                continue
            can_buy_votes.append(item.get("can_buy", False))
            amt = item.get("suggested_amount", 0)
            if isinstance(amt, (int, float)) and amt > 0:
                amounts.append(float(amt))
            for key in scores:
                val = item.get(f"{key}_score", 0)
                if isinstance(val, (int, float)) and val:
                    scores[key].append(float(val))
            if not first_text:
                first_text = item
            if item.get("reasoning"):
                reasonings.append(f"[{provider}] {item['reasoning']}")
            if item.get("trend_analysis"):
                trend_analyses.append(item["trend_analysis"])
            if item.get("entry_analysis"):
                entry_analyses.append(item["entry_analysis"])
            if item.get("holding_stock_analysis"):
                holding_stock_analyses.append(item["holding_stock_analysis"])
            if item.get("risk_warning"):
                risk_warnings.append(item["risk_warning"])

        # Majority vote for can_buy
        true_count = sum(1 for v in can_buy_votes if v)
        can_buy = true_count > len(can_buy_votes) / 2 if can_buy_votes else False
        agreement = true_count / len(can_buy_votes) if can_buy_votes else 0

        # Median suggested amount
        amounts.sort()
        median_amount = round(statistics.median(amounts)) if amounts else 0

        # Average scores
        avg_scores = {}
        for k, v in scores.items():
            avg_scores[f"{k}_score"] = round(sum(v) / len(v), 1) if v else 0

        consensus_funds[code] = {
            "code": code,
            "can_buy": can_buy,
            "ai_agreement": f"{agreement:.0%}",
            "ai_count": len(can_buy_votes),
            "suggested_amount": median_amount,
            "entry_strategy": first_text.get("entry_strategy", "--"),
            "entry_plan": first_text.get("entry_plan", "--"),
            "support_level": first_text.get("support_level", "--"),
            "resistance_level": first_text.get("resistance_level", "--"),
            "stop_loss": first_text.get("stop_loss", "--"),
            "target_price": first_text.get("target_price", "--"),
            "conditions_to_wait_for": first_text.get("conditions_to_wait_for", "--"),
            "technical_score": avg_scores.get("technical_score", 0),
            "manager_score": avg_scores.get("manager_score", 0),
            "allocation_score": avg_scores.get("allocation_score", 0),
            "scale_score": avg_scores.get("scale_score", 0),
            "overall_score": avg_scores.get("overall_score", 0),
            "trend_analysis": trend_analyses[0] if trend_analyses else "--",
            "entry_analysis": entry_analyses[0] if entry_analyses else "--",
            "holding_stock_analysis": holding_stock_analyses[0] if holding_stock_analyses else "--",
            "risk_warning": risk_warnings[0] if risk_warnings else "--",
            "reasoning": " | ".join(reasonings[:3]) if reasonings else "--",
        }

    return {
        "funds": consensus_funds,
        "market_overview": " ".join(market_overviews) if market_overviews else "",
        "watchlist_summary": " ".join(watchlist_summaries) if watchlist_summaries else "",
        "global_macro_analysis": " ".join(global_macro_analyses) if global_macro_analyses else "",
        "ai_count": len(parsed),
        "error": "",
    }


def run_multi_ai_analysis(label, codes, funds, indices, global_markets=None):
    """Send the same fund analysis prompt to all configured AIs in parallel.
    Returns consensus dict from cross_validate_ai_results.
    """
    if global_markets is None:
        global_markets = {}
    user_prompt = build_fund_context_for_ai(label, codes, funds, indices, global_markets)

    # Safety cap: truncate if excessively long (all models support 128K+)
    if len(user_prompt) > 100000:
        user_prompt = user_prompt[:100000] + "\n...[内容截断]"

    log(f"AI analysis prompt for {label}: {len(user_prompt)} chars")

    results = {}
    errors = []

    def call_one(config):
        return call_ai_api(config, AI_SYSTEM_PROMPT, user_prompt)

    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        futures = {}
        for name, config in AI_APIS.items():
            futures[executor.submit(call_one, config)] = name

        for future in concurrent.futures.as_completed(futures):
            name = futures[future]
            try:
                result = future.result()
                results[name] = result
                if result["error"]:
                    errors.append(f"{name}: {result['error']}")
                    log(f"AI [{name}] ERROR: {result['error'][:200]}")
                else:
                    log(f"AI [{name}] OK ({len(result['content'])} chars)")
            except Exception as e:
                errors.append(f"{name}: {str(e)}")
                log(f"AI [{name}] EXCEPTION: {e}")

    log(
        f"AI analysis for {label}: {len(results) - len(errors)}/{len(AI_APIS)} succeeded, "
        f"{len(errors)} errors"
    )

    consensus = cross_validate_ai_results(results, codes)
    consensus["errors"] = errors
    consensus["raw_results"] = results  # preserved for debugging

    fund_count = len(consensus.get("funds", {}))
    log(f"AI consensus for {label}: {fund_count}/{len(codes)} funds have recommendations")
    return consensus


def run_watchlist_ai_analysis(codes, funds, indices, global_markets=None):
    """Send watchlist fund analysis prompt to all configured AIs in parallel.
    Returns consensus dict from cross_validate_watchlist_results.
    """
    if global_markets is None:
        global_markets = {}
    user_prompt = build_watchlist_context_for_ai(codes, funds, indices, global_markets)

    # Safety cap: truncate if excessively long
    if len(user_prompt) > 100000:
        user_prompt = user_prompt[:100000] + "\n...[内容截断]"

    log(f"Watchlist AI analysis prompt: {len(user_prompt)} chars")

    results = {}
    errors = []

    def call_one(config):
        return call_ai_api(config, WATCHLIST_SYSTEM_PROMPT, user_prompt)

    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        futures = {}
        for name, config in AI_APIS.items():
            futures[executor.submit(call_one, config)] = name

        for future in concurrent.futures.as_completed(futures):
            name = futures[future]
            try:
                result = future.result()
                results[name] = result
                if result["error"]:
                    errors.append(f"{name}: {result['error']}")
                    log(f"Watchlist AI [{name}] ERROR: {result['error'][:200]}")
                else:
                    log(f"Watchlist AI [{name}] OK ({len(result['content'])} chars)")
            except Exception as e:
                errors.append(f"{name}: {str(e)}")
                log(f"Watchlist AI [{name}] EXCEPTION: {e}")

    log(
        f"Watchlist AI analysis: {len(results) - len(errors)}/{len(AI_APIS)} succeeded, "
        f"{len(errors)} errors"
    )

    consensus = cross_validate_watchlist_results(results, codes)
    consensus["errors"] = errors
    consensus["raw_results"] = results  # preserved for debugging

    fund_count = len(consensus.get("funds", {}))
    log(f"Watchlist AI consensus: {fund_count}/{len(codes)} funds have recommendations")
    return consensus


# =============================================================================
# SECTION F: Report Generation — Shared Helpers
# =============================================================================


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False, font_size=8.5):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(font_size)
            run.bold = bold


def style_table(table):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    style_table(table)
    document.add_paragraph()
    return table


def setup_document(title, subtitle):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)
    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_paragraph.add_run(subtitle)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 112, 133)
    return document


def fund_rows(codes, funds):
    rows = []
    for code in codes:
        fund = funds[code]
        latest = fund.get("latest") or {}
        rows.append(
            [
                code,
                fund["name"],
                latest.get("date", "--"),
                latest.get("nav", "--"),
                pct_text(latest.get("day_return")),
                pct_text(fund["m1"]),
                pct_text(fund["m3"]),
                pct_text(fund["m6"]),
                pct_text(fund["y1"]),
                fund["manager"],
                fund["scale"],
                fund["allocation"],
            ]
        )
    return rows


# =============================================================================
# SECTION G: Word Report Generation (enhanced with AI analysis)
# =============================================================================


def add_fund_analysis_section(doc, codes, funds, ai_consensus):
    """Add per-fund deep analysis section to the Word document."""
    doc.add_heading("四、单只基金深度分析（多AI共识）", level=1)

    ai_count = ai_consensus.get("ai_count", 0)
    if ai_count == 0:
        doc.add_paragraph(
            "⚠ AI分析未能完成（所有AI接口调用失败或返回格式异常）。"
            "以下为基于公开数据的静态分析框架。"
        )
    else:
        meta_note = []
        if ai_consensus.get("global_macro_analysis"):
            meta_note.append(f"🌍 全球宏观分析：{ai_consensus['global_macro_analysis']}")
        if ai_consensus.get("market_overview"):
            meta_note.append(f"市场判断：{ai_consensus['market_overview']}")
        if ai_consensus.get("portfolio_analysis"):
            meta_note.append(f"组合分析：{ai_consensus['portfolio_analysis']}")
        if meta_note:
            for note in meta_note:
                doc.add_paragraph(note)
        doc.add_paragraph(
            f"（以上分析由 {ai_count} 个AI模型共同生成，经交叉验证后取共识结果）"
        )

    # === Position recommendation summary table ===
    doc.add_heading("4.1 仓位操作建议汇总", level=2)
    position_headers = [
        "代码", "基金名称", "综合评分", "操作建议", "金额变动",
        "百分比", "AI一致率", "主要风险"
    ]
    position_rows = []
    for code in codes:
        cf = ai_consensus.get("funds", {}).get(code, {})
        fund = funds[code]
        action = cf.get("action", "--")
        pct = cf.get("percentage_change", 0)
        amount = cf.get("amount_change", 0)

        if action == "增持" and amount:
            action_display = f"增持 +{amount}元"
        elif action == "减持" and amount:
            action_display = f"减持 -{amount}元"
        else:
            action_display = action if action else "--"

        pct_display = (
            f"+{pct}%" if action == "增持" and pct
            else (f"-{pct}%" if action == "减持" and pct else "--")
        )

        position_rows.append(
            [
                code,
                fund["name"],
                str(cf.get("overall_score", "--")),
                action_display,
                f"{amount}元" if amount else "--",
                pct_display,
                cf.get("ai_agreement", "--"),
                (cf.get("risk_warning", "--") or "--")[:80],
            ]
        )
    add_table(doc, position_headers, position_rows)

    # === Per-fund detailed analysis ===
    doc.add_heading("4.2 逐只基金详细分析", level=2)
    for code in codes:
        cf = ai_consensus.get("funds", {}).get(code, {})
        fund = funds[code]

        doc.add_heading(f"{code} {fund['name']}", level=3)

        if fund.get("error"):
            doc.add_paragraph(f"数据获取异常：{fund['error']}")
            doc.add_paragraph("因数据不足，无法进行AI深度分析，建议维持当前仓位。")
            continue

        # Score breakdown mini-table
        score_headers = ["技术面", "基金经理", "行业配置", "规模", "综合"]
        score_rows = [
            [
                str(cf.get("technical_score", "--")),
                str(cf.get("manager_score", "--")),
                str(cf.get("allocation_score", "--")),
                str(cf.get("scale_score", "--")),
                str(cf.get("overall_score", "--")),
            ]
        ]
        add_table(doc, score_headers, score_rows)

        # Action & amount
        action = cf.get("action", "--")
        amount = cf.get("amount_change", 0)
        pct = cf.get("percentage_change", 0)
        action_paragraph = doc.add_paragraph()
        action_run = action_paragraph.add_run(f"操作建议：{action}")
        action_run.bold = True
        if action in ("增持", "减持") and amount:
            direction = "+" if action == "增持" else "-"
            action_paragraph.add_run(
                f"  {direction}{amount}元（{direction}{pct}%），"
                f"AI一致率：{cf.get('ai_agreement', '--')}"
            )

        # Trend analysis
        trend_text = cf.get("trend_analysis", "--")
        if trend_text and trend_text != "--":
            doc.add_paragraph(f"趋势分析：{trend_text}")

        # Holding stock analysis (NEW)
        holding_text = cf.get("holding_stock_analysis", "--")
        if holding_text and holding_text != "--":
            holding_para = doc.add_paragraph()
            holding_run = holding_para.add_run(f"📊 重仓股分析：{holding_text}")
            holding_run.font.color.rgb = RGBColor(31, 78, 121)

        # Reasoning
        reasoning_text = cf.get("reasoning", "--")
        if reasoning_text and reasoning_text != "--":
            doc.add_paragraph(f"分析理由：{reasoning_text}")

        # Risk warning
        risk_text = cf.get("risk_warning", "--")
        if risk_text and risk_text != "--":
            risk_paragraph = doc.add_paragraph()
            risk_run = risk_paragraph.add_run(f"⚠ 风险提示：{risk_text}")
            risk_run.font.color.rgb = RGBColor(180, 35, 24)

        # Fund basic data summary
        ks = fund.get("kline_summary", {})
        if ks:
            doc.add_paragraph(
                f"技术数据：近{ks.get('period_days', '--')}日趋势{ks.get('recent_trend', '--')}，"
                f"区间收益{ks.get('period_return', '--')}，"
                f"波动率{ks.get('volatility', '--')}"
            )


def add_watchlist_analysis_section(doc, codes, funds, wl_consensus):
    """Add watchlist fund entry analysis section to the Word document."""
    doc.add_heading("八、自选观察基金买入分析（多AI共识）", level=1)

    ai_count = wl_consensus.get("ai_count", 0)
    if ai_count == 0:
        doc.add_paragraph("⚠ 自选观察AI分析未能完成（所有AI接口调用失败或返回格式异常）。")
        return

    # Market overview and summary
    if wl_consensus.get("market_overview"):
        doc.add_paragraph(f"市场判断：{wl_consensus['market_overview']}")
    if wl_consensus.get("watchlist_summary"):
        doc.add_paragraph(f"整体建议：{wl_consensus['watchlist_summary']}")
    doc.add_paragraph(
        f"（以上分析由 {ai_count} 个AI模型共同生成，经交叉验证后取共识结果）"
    )

    # === Buy recommendation summary table ===
    doc.add_heading("8.1 买入建议汇总", level=2)
    wl_headers = [
        "代码", "基金名称", "综合评分", "是否买入", "建议金额",
        "买入策略", "AI一致率", "主要风险"
    ]
    wl_rows = []
    for code in codes:
        cf = wl_consensus.get("funds", {}).get(code, {})
        fund = funds[code]
        can_buy = cf.get("can_buy", False)
        buy_display = "✅ 建议买入" if can_buy else "❌ 暂不建议"
        amount = cf.get("suggested_amount", 0)
        wl_rows.append(
            [
                code,
                fund["name"],
                str(cf.get("overall_score", "--")),
                buy_display,
                f"{amount}元" if amount else "--",
                cf.get("entry_strategy", "--"),
                cf.get("ai_agreement", "--"),
                (cf.get("risk_warning", "--") or "--")[:80],
            ]
        )
    add_table(doc, wl_headers, wl_rows)

    # === Per-fund detailed analysis ===
    doc.add_heading("8.2 逐只基金详细分析", level=2)
    for code in codes:
        cf = wl_consensus.get("funds", {}).get(code, {})
        fund = funds[code]

        doc.add_heading(f"{code} {fund['name']}", level=3)

        if fund.get("error"):
            doc.add_paragraph(f"数据获取异常：{fund['error']}")
            doc.add_paragraph("因数据不足，无法进行AI买入分析。")
            continue

        # Score breakdown mini-table
        score_headers = ["技术面", "基金经理", "行业配置", "规模", "综合"]
        score_rows = [
            [
                str(cf.get("technical_score", "--")),
                str(cf.get("manager_score", "--")),
                str(cf.get("allocation_score", "--")),
                str(cf.get("scale_score", "--")),
                str(cf.get("overall_score", "--")),
            ]
        ]
        add_table(doc, score_headers, score_rows)

        # Buy decision
        can_buy = cf.get("can_buy", False)
        decision_para = doc.add_paragraph()
        decision_run = decision_para.add_run(
            f"买入判断：{'✅ 建议买入' if can_buy else '❌ 暂不建议买入'}"
        )
        decision_run.bold = True
        if can_buy:
            decision_para.add_run(
                f"  建议金额{cf.get('suggested_amount', 0)}元，"
                f"策略：{cf.get('entry_strategy', '--')}"
            )

        # Trend and entry analysis
        trend_text = cf.get("trend_analysis", "--")
        if trend_text and trend_text != "--":
            doc.add_paragraph(f"趋势分析：{trend_text}")
        entry_text = cf.get("entry_analysis", "--")
        if entry_text and entry_text != "--":
            doc.add_paragraph(f"入场分析：{entry_text}")

        # Holding stock analysis (NEW)
        holding_text = cf.get("holding_stock_analysis", "--")
        if holding_text and holding_text != "--":
            holding_para = doc.add_paragraph()
            holding_run = holding_para.add_run(f"📊 重仓股分析：{holding_text}")
            holding_run.font.color.rgb = RGBColor(31, 78, 121)

        # Support / Resistance
        support = cf.get("support_level", "--")
        resistance = cf.get("resistance_level", "--")
        if support != "--" or resistance != "--":
            doc.add_paragraph(f"支撑位：{support}  |  压力位：{resistance}")

        # Entry plan
        entry_plan = cf.get("entry_plan", "--")
        if entry_plan and entry_plan != "--":
            doc.add_paragraph(f"买入计划：{entry_plan}")

        # Stop loss and target
        stop_loss = cf.get("stop_loss", "--")
        if stop_loss and stop_loss != "--":
            doc.add_paragraph(f"止损条件：{stop_loss}")
        target = cf.get("target_price", "--")
        if target and target != "--":
            doc.add_paragraph(f"目标价位：{target}")

        # Conditions to wait for (if can't buy)
        conditions = cf.get("conditions_to_wait_for", "--")
        if not can_buy and conditions and conditions != "--":
            cond_para = doc.add_paragraph()
            cond_run = cond_para.add_run(f"等待条件：{conditions}")
            cond_run.font.color.rgb = RGBColor(180, 35, 24)

        # Reasoning
        reasoning_text = cf.get("reasoning", "--")
        if reasoning_text and reasoning_text != "--":
            doc.add_paragraph(f"分析理由：{reasoning_text}")

        # Risk warning
        risk_text = cf.get("risk_warning", "--")
        if risk_text and risk_text != "--":
            risk_para = doc.add_paragraph()
            risk_run = risk_para.add_run(f"⚠ 风险提示：{risk_text}")
            risk_run.font.color.rgb = RGBColor(180, 35, 24)

        # Technical data summary
        ks = fund.get("kline_summary", {})
        if ks:
            doc.add_paragraph(
                f"技术数据：近{ks.get('period_days', '--')}日趋势{ks.get('recent_trend', '--')}，"
                f"区间收益{ks.get('period_return', '--')}，"
                f"波动率{ks.get('volatility', '--')}"
            )


def make_docx(label, codes, funds, indices, output_path, ai_consensus=None, wl_consensus=None):
    if ai_consensus is None:
        ai_consensus = {}
    if wl_consensus is None:
        wl_consensus = {}

    stats = group_stats(codes, funds)
    doc = setup_document(
        f"{label}每日复盘与仓位建议",
        f"运行日期：{RUN_DATE.isoformat()}    分析日期：{ANALYSIS_DATE.isoformat()}"
        f"    数据源：东方财富/天天基金公开数据    分析引擎：ChatGPT + Claude + DeepSeek 多AI共识",
    )

    # --- Section 1: Summary ---
    doc.add_heading("一、摘要", level=1)
    add_table(
        doc,
        ["组合", "有数据基金数", "平均日涨跌", "上涨只数", "下跌只数"],
        [[label, stats["count"], stats["avg"], stats["up"], stats["down"]]],
    )
    if ai_consensus.get("portfolio_recommendation"):
        doc.add_paragraph(f"AI组合建议：{ai_consensus['portfolio_recommendation']}")
    else:
        doc.add_paragraph(
            "结论：昨日市场偏向成长和高弹性资产。若组合已明显盈利，"
            "不建议因单日上涨追高；更适合检查集中度、重复暴露和单只基金权重。"
        )

    # --- Section 2: Market Environment ---
    doc.add_heading("二、市场环境", level=1)
    add_table(
        doc,
        ["指数", "日期", "收盘", "涨跌幅", "涨跌额", "日内区间"],
        [
            [
                item["name"],
                item["date"],
                item["close"],
                pct_text(item["pct"]),
                item["chg"],
                f"{item['low']} - {item['high']}",
            ]
            for item in indices
        ],
    )

    # --- Section 3: Fund Details ---
    doc.add_heading("三、基金明细", level=1)
    add_table(
        doc,
        [
            "代码", "基金名称", "净值日期", "单位净值", "日涨跌",
            "近1月", "近3月", "近6月", "近1年",
            "基金经理", "规模", "资产配置",
        ],
        fund_rows(codes, funds),
    )

    # --- Section 4: Per-Fund AI Deep Analysis (NEW) ---
    add_fund_analysis_section(doc, codes, funds, ai_consensus)

    # --- Section 5: Portfolio Analysis (renumbered from 4) ---
    doc.add_heading("五、组合分析", level=1)
    doc.add_paragraph(
        "1. 成长/科技相关基金在昨日市场环境中更容易受益，"
        "但短期涨幅越集中，后续波动和回撤风险也越高。"
    )
    doc.add_paragraph(
        "2. 两个组合中存在重复代码 021528 和 022365，"
        "需要在总账户层面合并计算暴露，避免实际权重高于预期。"
    )
    doc.add_paragraph(
        "3. QDII、港股、黄金等资产净值可能滞后一个交易日，"
        "判断组合表现时应区分披露日期。"
    )
    doc.add_paragraph(
        "4. AI分析基于近期K线走势、经理评估、行业配置等多维度数据，"
        "结合三模型交叉验证给出仓位建议，但无法预测突发政策或市场事件。"
    )

    # --- Section 6: Position Management Framework ---
    doc.add_heading("六、仓位管理框架", level=1)
    add_table(
        doc,
        ["场景", "建议动作"],
        [
            [
                "AI建议增持",
                "参考AI分析的具体百分比，结合自身资金情况分批加仓，不一次性满仓。",
            ],
            [
                "AI建议减持",
                "考虑分批止盈或降低单只权重，优先减持评分最低、风险最高的基金。",
            ],
            [
                "AI建议持有",
                "维持当前仓位观望，等待更明确的趋势信号或基本面变化再做决策。",
            ],
            [
                "已有明显盈利",
                "不因单日上涨追高加仓，优先检查组合集中度和单只基金权重。",
            ],
            [
                "同风格基金重复较多",
                "合并计算科技、成长、半导体、港股等主题暴露，控制总权重。",
            ],
            [
                "仍在定投",
                "维持小额、分批、纪律化执行；只有回撤到预设区间且基本面未恶化时再提高定投金额。",
            ],
        ],
    )

    # --- Section 7: Risk Warning & Data Sources ---
    doc.add_heading("七、风险提示与数据来源", level=1)
    doc.add_paragraph(
        "本报告基于公开数据与AI模型分析自动生成，不构成个性化投资顾问意见，"
        "不承诺收益。AI分析可能受数据质量、模型偏差等因素影响。"
        "最终操作应结合自身风险承受能力、资金期限和实际持仓成本独立判断。"
    )
    add_table(
        doc,
        ["基金代码", "来源链接", "异常"],
        [
            [code, funds[code]["source"], funds[code]["error"] or "--"]
            for code in codes
        ],
    )

    # --- Section 八: Watchlist Entry Analysis (NEW) ---
    if wl_consensus and wl_consensus.get("ai_count", 0) > 0:
        add_watchlist_analysis_section(doc, WATCHLIST_CODES, funds, wl_consensus)

    doc.save(output_path)


# =============================================================================
# SECTION H: HTML Email Generation (enhanced with AI analysis)
# =============================================================================


def html_pct(value):
    number = to_float(value)
    if number is None:
        return "--"
    color = "#b42318" if number > 0 else ("#067647" if number < 0 else "#475467")
    return f'<span style="color:{color};font-weight:600;">{number:+.2f}%</span>'


def html_table(headers, rows):
    header_html = "".join(
        f"<th>{html.escape(str(header))}</th>" for header in headers
    )
    row_html = []
    for row in rows:
        row_html.append(
            "<tr>"
            + "".join(f"<td>{cell}</td>" for cell in row)
            + "</tr>"
        )
    return (
        "<table><thead><tr>"
        + header_html
        + "</tr></thead><tbody>"
        + "".join(row_html)
        + "</tbody></table>"
    )


def fund_analysis_html(codes, funds, ai_consensus):
    """Generate HTML for the per-fund AI analysis section."""
    ai_count = ai_consensus.get("ai_count", 0)

    if ai_count == 0:
        return """
        <h2>AI深度分析</h2>
        <div class="note" style="background:#fff3cd;padding:12px;border-radius:6px;margin:12px 0;">
          ⚠ AI分析未能完成（所有AI接口调用失败或返回格式异常）。请检查API密钥配置和网络连接。
        </div>"""

    parts = []

    # Global macro analysis (NEW)
    global_macro = ai_consensus.get("global_macro_analysis", "")
    if global_macro:
        parts.append(
            f'<div style="background:#eef2ff;padding:12px;border-radius:6px;margin:12px 0;border-left:4px solid:#6366f1;">'
            f'<strong>🌍 全球宏观分析</strong><br>{html.escape(global_macro)}</div>'
        )

    # Market overview from AI
    market_overview = ai_consensus.get("market_overview", "")
    if market_overview:
        parts.append(
            f'<div style="background:#f0f7ff;padding:12px;border-radius:6px;margin:12px 0;border-left:4px solid:#2e90fa;">'
            f'<strong>AI市场判断</strong><br>{html.escape(market_overview)}</div>'
        )

    portfolio_analysis = ai_consensus.get("portfolio_analysis", "")
    if portfolio_analysis:
        parts.append(
            f'<div style="background:#f9fafb;padding:12px;border-radius:6px;margin:12px 0;">'
            f'<strong>组合分析</strong><br>{html.escape(portfolio_analysis)}</div>'
        )

    # Position recommendation table
    position_rows = []
    for code in codes:
        cf = ai_consensus.get("funds", {}).get(code, {})
        fund = funds[code]
        action = cf.get("action", "--")
        pct = cf.get("percentage_change", 0)
        amount = cf.get("amount_change", 0)

        # Action badge
        if action == "增持":
            action_color = "#b42318"
            action_bg = "#fef3f2"
        elif action == "减持":
            action_color = "#067647"
            action_bg = "#ecfdf3"
        else:
            action_color = "#475467"
            action_bg = "#f9fafb"

        action_html = (
            f'<span style="display:inline-block;background:{action_bg};color:{action_color};'
            f'font-weight:700;padding:4px 10px;border-radius:4px;">{action}</span>'
        )
        if action in ("增持", "减持") and amount:
            sign = "+" if action == "增持" else "-"
            action_html += (
                f'<br><span style="font-size:11px;color:{action_color};">'
                f'{sign}{amount}元 ({sign}{pct}%)</span>'
            )

        # Score bar
        score = cf.get("overall_score", 0)
        if score >= 7:
            bar_color = "#2e90fa"
        elif score >= 4:
            bar_color = "#f79009"
        else:
            bar_color = "#b42318"
        score_html = (
            f'<div style="display:flex;align-items:center;gap:6px;">'
            f'<div style="background:#f0f0f0;border-radius:3px;width:50px;height:6px;">'
            f'<div style="background:{bar_color};height:6px;border-radius:3px;width:{min(score * 10, 100)}%;"></div>'
            f'</div><span style="font-weight:600;">{score}</span></div>'
        )

        risk_text = cf.get("risk_warning", "--") or "--"
        position_rows.append(
            [
                html.escape(code),
                html.escape(fund["name"]),
                score_html,
                action_html,
                html.escape(cf.get("ai_agreement", "--")),
                html.escape(risk_text[:80] + ("..." if len(risk_text) > 80 else "")),
            ]
        )

    parts.append(
        f'<div style="margin:12px 0;font-size:12px;color:#667085;">'
        f'AI共识仓位建议（由 {ai_count} 个AI模型交叉验证）</div>'
    )
    parts.append(
        html_table(
            ["代码", "名称", "评分", "操作建议", "AI一致率", "主要风险"],
            position_rows,
        )
    )

    # Portfolio-level recommendation
    portfolio_rec = ai_consensus.get("portfolio_recommendation", "")
    if portfolio_rec:
        parts.append(
            f'<div style="background:#f9fafb;padding:12px;border-radius:6px;margin:12px 0;">'
            f'<strong>组合层面建议</strong><br>{html.escape(portfolio_rec)}</div>'
        )

    return "<h2>AI深度分析与仓位建议</h2>\n" + "\n".join(parts)


def watchlist_analysis_html(codes, funds, wl_consensus):
    """Generate HTML for the watchlist fund entry analysis section."""
    ai_count = wl_consensus.get("ai_count", 0)

    if ai_count == 0:
        return ""

    parts = []

    # Global macro analysis (NEW)
    global_macro = wl_consensus.get("global_macro_analysis", "")
    if global_macro:
        parts.append(
            f'<div style="background:#eef2ff;padding:12px;border-radius:6px;margin:12px 0;border-left:4px solid:#6366f1;">'
            f'<strong>🌍 全球宏观分析</strong><br>{html.escape(global_macro)}</div>'
        )

    # Market overview
    market_overview = wl_consensus.get("market_overview", "")
    if market_overview:
        parts.append(
            f'<div style="background:#f0f7ff;padding:12px;border-radius:6px;margin:12px 0;border-left:4px solid:#2e90fa;">'
            f'<strong>AI市场判断（自选观察）</strong><br>{html.escape(market_overview)}</div>'
        )

    # Watchlist summary
    watchlist_summary = wl_consensus.get("watchlist_summary", "")
    if watchlist_summary:
        parts.append(
            f'<div style="background:#f9fafb;padding:12px;border-radius:6px;margin:12px 0;">'
            f'<strong>整体观察建议</strong><br>{html.escape(watchlist_summary)}</div>'
        )

    # Summary table
    wl_rows = []
    for code in codes:
        cf = wl_consensus.get("funds", {}).get(code, {})
        fund = funds[code]
        can_buy = cf.get("can_buy", False)

        if can_buy:
            action_color = "#b42318"
            action_bg = "#fef3f2"
            buy_text = "✅ 建议买入"
        else:
            action_color = "#067647"
            action_bg = "#ecfdf3"
            buy_text = "❌ 暂不建议"

        action_html = (
            f'<span style="display:inline-block;background:{action_bg};color:{action_color};'
            f'font-weight:700;padding:4px 10px;border-radius:4px;">{buy_text}</span>'
        )
        amount = cf.get("suggested_amount", 0)
        if can_buy and amount:
            sign = "+"
            action_html += (
                f'<br><span style="font-size:11px;color:{action_color};">'
                f'{sign}{amount}元</span>'
            )

        # Score bar
        score = cf.get("overall_score", 0)
        if score >= 7:
            bar_color = "#2e90fa"
        elif score >= 4:
            bar_color = "#f79009"
        else:
            bar_color = "#b42318"
        score_html = (
            f'<div style="display:flex;align-items:center;gap:6px;">'
            f'<div style="background:#f0f0f0;border-radius:3px;width:50px;height:6px;">'
            f'<div style="background:{bar_color};height:6px;border-radius:3px;width:{min(score * 10, 100)}%;"></div>'
            f'</div><span style="font-weight:600;">{score}</span></div>'
        )

        risk_text = cf.get("risk_warning", "--") or "--"
        wl_rows.append(
            [
                html.escape(code),
                html.escape(fund["name"]),
                score_html,
                action_html,
                html.escape(cf.get("ai_agreement", "--")),
                html.escape(risk_text[:80] + ("..." if len(risk_text) > 80 else "")),
            ]
        )

    parts.append(
        f'<div style="margin:12px 0;font-size:12px;color:#667085;">'
        f'自选观察基金买入建议（由 {ai_count} 个AI模型交叉验证）</div>'
    )
    parts.append(
        html_table(
            ["代码", "名称", "评分", "买入建议", "AI一致率", "主要风险"],
            wl_rows,
        )
    )

    return "<h2>自选观察基金买入分析</h2>\n" + "\n".join(parts)


def email_body(label, codes, funds, indices, ai_consensus=None, wl_consensus=None):
    if ai_consensus is None:
        ai_consensus = {}
    if wl_consensus is None:
        wl_consensus = {}

    stats = group_stats(codes, funds)

    # Market index rows
    index_rows = [
        [
            html.escape(item["name"]),
            html.escape(item["date"]),
            html.escape(str(item["close"])),
            html_pct(item["pct"]),
            html.escape(str(item["chg"])),
            html.escape(f"{item['low']} - {item['high']}"),
        ]
        for item in indices
    ]

    # Fund detail rows
    fund_rows_html = []
    for code in codes:
        fund = funds[code]
        latest = fund.get("latest") or {}
        fund_rows_html.append(
            [
                html.escape(code),
                html.escape(fund["name"]),
                html.escape(str(latest.get("date", "--"))),
                html.escape(str(latest.get("nav", "--"))),
                html_pct(latest.get("day_return")),
                html_pct(fund["m1"]),
                html_pct(fund["m3"]),
                html_pct(fund["m6"]),
                html_pct(fund["y1"]),
            ]
        )

    css = """
    body { font-family: Arial, 'Microsoft YaHei', sans-serif; color:#101828; line-height:1.55; }
    h1 { font-size:20px; margin:0 0 6px; }
    h2 { font-size:16px; margin:22px 0 8px; border-left:4px solid #2e90fa; padding-left:8px; }
    .meta,.note { color:#667085; font-size:13px; }
    table { border-collapse:collapse; width:100%; font-size:13px; margin:8px 0 14px; }
    th { background:#f2f4f7; border:1px solid #d0d5dd; padding:7px 8px; text-align:left; white-space:nowrap; }
    td { border:1px solid #d0d5dd; padding:7px 8px; vertical-align:top; }
    .footer { color:#667085; font-size:12px; margin-top:24px; border-top:1px solid #d0d5dd; padding-top:12px; }
    """

    # Build AI analysis section
    ai_section = fund_analysis_html(codes, funds, ai_consensus)
    wl_section = watchlist_analysis_html(WATCHLIST_CODES, funds, wl_consensus)

    return f"""<!doctype html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>{label}每日复盘与仓位建议</h1>
<div class="meta">运行日期：{RUN_DATE.isoformat()} ｜ 分析日期：{ANALYSIS_DATE.isoformat()} ｜ 数据源：东方财富/天天基金公开数据 ｜ 分析引擎：ChatGPT + Claude + DeepSeek 多AI共识</div>

<h2>组合表现</h2>
{html_table(["组合", "有数据基金数", "平均日涨跌", "上涨只数", "下跌只数"], [[label, stats["count"], stats["avg"], stats["up"], stats["down"]]])}

<h2>市场概况</h2>
{html_table(["指数", "日期", "收盘", "涨跌幅", "涨跌额", "日内区间"], index_rows)}

<h2>基金明细</h2>
{html_table(["代码", "基金名称", "净值日期", "单位净值", "日涨跌", "近1月", "近3月", "近6月", "近1年"], fund_rows_html)}

{ai_section}

{wl_section}

<h2>仓位管理框架</h2>
{html_table(["场景", "建议动作"], [
    ["AI建议增持", "参考AI分析的具体百分比，结合自身资金情况分批加仓，不一次性满仓。"],
    ["AI建议减持", "考虑分批止盈或降低单只权重，优先减持评分最低、风险最高的基金。"],
    ["AI建议持有", "维持当前仓位观望，等待更明确的趋势信号或基本面变化再做决策。"],
    ["组合再平衡", "以AI评分为参考，定期检查集中度和风格暴露，避免重复持仓。"],
    ["仍在定投", "维持小额、分批、纪律化执行；只有回撤到预设区间且基本面未恶化时再提高定投金额。"],
])}

<div class="footer">
<p>本邮件为公开数据+AI模型分析自动生成，不构成个性化投资建议或收益承诺。AI分析可能受数据质量、模型偏差等因素影响。最终操作应结合自身风险承受能力独立判断。完整表格见附件 Word 文档。</p>
<p>分析引擎：OpenAI ChatGPT (gpt-4o) + Anthropic Claude (claude-sonnet-4) + DeepSeek (deepseek-chat) | 数据源：东方财富</p>
</div>
</body></html>"""


# =============================================================================
# SECTION I: Email Sending
# =============================================================================


def send_email(label, recipient, docx_path, codes, funds, indices, ai_consensus=None, wl_consensus=None):
    if ai_consensus is None:
        ai_consensus = {}
    if wl_consensus is None:
        wl_consensus = {}

    user = user_env("FUND_REPORT_SMTP_USER") or "your_gmail@gmail.com"
    password = user_env("FUND_REPORT_SMTP_PASS")
    if not password:
        raise RuntimeError("FUND_REPORT_SMTP_PASS is missing")

    msg = EmailMessage()
    msg["From"] = user
    msg["To"] = recipient
    msg["Subject"] = (
        f"{label}每日复盘与仓位建议 - {ANALYSIS_DATE.isoformat()}"
    )
    msg.set_content(
        f"{label}每日复盘与仓位建议已生成。请查看附件 Word 文档。"
        f"\n\n本报告由 ChatGPT + Claude + DeepSeek 三AI模型共同分析生成。"
    )
    msg.add_alternative(
        email_body(label, codes, funds, indices, ai_consensus, wl_consensus), subtype="html"
    )
    data = docx_path.read_bytes()
    msg.add_attachment(
        data,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=docx_path.name,
    )

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=45) as smtp:
        smtp.ehlo()
        smtp.starttls()
        smtp.ehlo()
        smtp.login(user, password)
        smtp.send_message(msg)


# =============================================================================
# SECTION J: Orchestration
# =============================================================================


def run():
    RUN_DIR.mkdir(parents=True, exist_ok=True)
    beijing_hour = beijing_now().hour
    is_evening = beijing_hour >= 15
    log(
        f"starting enhanced report run; run_date={RUN_DATE.isoformat()} "
        f"analysis_date={ANALYSIS_DATE.isoformat()} "
        f"beijing_hour={beijing_hour} is_evening={is_evening}"
    )

    # --- Phase 1: Fetch market indices ---
    log("Phase 1/5: Fetching market indices...")
    indices = [
        fetch_index("上证指数", "1.000001"),
        fetch_index("深证成指", "0.399001"),
        fetch_index("创业板指", "0.399006"),
        fetch_index("科创50", "1.000688"),
        fetch_index("沪深300", "1.000300"),
        fetch_index("恒生指数", "100.HSI"),
    ]
    log(f"Fetched {len(indices)} indices")
    # Fetch global market data (US indices) for macro context
    global_markets = fetch_global_markets()
    log(f"Global markets: {global_markets.get('macro_signals', 'N/A')[:120]}")

    # --- Phase 2: Fetch all fund data (20 unique funds including watchlist) ---
    log("Phase 2/5: Fetching fund data (with watchlist funds)...")
    funds = {}
    all_codes = sorted(set(FUND_ONE + FUND_TWO + WATCHLIST_CODES))
    for code in all_codes:
        funds[code] = fetch_fund(code)
        error_flag = "ERR" if funds[code]["error"] else "OK"
        kline_days = len(funds[code].get("kline_summary", {}))
        log(
            f"  fund {code} [{error_flag}]: {funds[code]['name']} "
            f"(kline_days={kline_days}) error={funds[code]['error'] or '-'}"
        )

    # --- Phase 3: Multi-AI analysis for portfolios ---
    log("Phase 3/5: Running multi-AI deep analysis for portfolios...")
    log("  Starting AI analysis for 基金一...")
    ai_one = run_multi_ai_analysis("基金一", FUND_ONE, funds, indices, global_markets)

    log("  Starting AI analysis for 基金二...")
    ai_two = run_multi_ai_analysis("基金二", FUND_TWO, funds, indices, global_markets)

    ai_one_consensus = ai_one if ai_one else {}
    ai_two_consensus = ai_two if ai_two else {}

    # --- Phase 3b: Watchlist AI analysis (evening only) ---
    wl_consensus = {}
    if is_evening:
        log("Phase 3b/5: Running watchlist AI analysis (evening run)...")
        wl_consensus = run_watchlist_ai_analysis(WATCHLIST_CODES, funds, indices, global_markets)
    else:
        log("Phase 3b/5: Skipping watchlist AI analysis (morning run, save API cost)")

    # --- Phase 4: Generate reports ---
    log("Phase 4/5: Generating reports...")
    doc1 = RUN_DIR / f"基金一_每日复盘_{ANALYSIS_DATE.isoformat()}.docx"
    doc2 = RUN_DIR / f"基金二_每日复盘_{ANALYSIS_DATE.isoformat()}.docx"

    make_docx("基金一", FUND_ONE, funds, indices, doc1, ai_one_consensus, wl_consensus)
    log(f"Created docx: {doc1}")
    make_docx("基金二", FUND_TWO, funds, indices, doc2, ai_two_consensus, wl_consensus)
    log(f"Created docx: {doc2}")

    # --- Phase 5: Send emails (evening only) ---
    if is_evening:
        log("Phase 5/5: Sending emails (evening run)...")
        send_email(
            "基金一", RECIPIENTS["基金一"], doc1, FUND_ONE, funds, indices,
            ai_one_consensus, wl_consensus,
        )
        log(f"Sent email for 基金一 to {RECIPIENTS['基金一']}")
        send_email(
            "基金二", RECIPIENTS["基金二"], doc2, FUND_TWO, funds, indices,
            ai_two_consensus, wl_consensus,
        )
        log(f"Sent email for 基金二 to {RECIPIENTS['基金二']}")
    else:
        log("Phase 5/5: Skipping email sending (before 15:00 Beijing time, emails only sent in evening run)")

    # Summary
    ai_one_count = len(ai_one_consensus.get("funds", {}))
    ai_two_count = len(ai_two_consensus.get("funds", {}))
    wl_count = len(wl_consensus.get("funds", {}))
    log(
        f"Report run completed. AI consensus: 基金一={ai_one_count}/{len(FUND_ONE)} funds, "
        f"基金二={ai_two_count}/{len(FUND_TWO)} funds, "
        f"自选观察={wl_count}/{len(WATCHLIST_CODES)} funds"
    )


# =============================================================================
# SECTION K: Main Entry Point
# =============================================================================

if __name__ == "__main__":
    try:
        run()
    except Exception:
        log("FAILED")
        log(traceback.format_exc())
        sys.exit(1)
